INPUT_SPEC = {
    "path": "string"
}

import os
import sys

BASE_PATH = os.path.abspath("E:/MutesHand")


def run(path):
    """
    Extract plain text from a local DOCX file.

    Returns structured dict for all cases.
    """
    try:
        _project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
        if _project_root not in sys.path:
            sys.path.insert(0, _project_root)
        from system.security.path_validator import validate_path

        validation = validate_path(path, BASE_PATH)
        if validation.get("status") == "failure":
            return validation

        full_path = validation["resolved_path"]

        # Check if file exists
        if not os.path.exists(full_path):
            return {"status": "failure", "reason": "file_not_found"}

        # Check extension
        if not full_path.lower().endswith(".docx"):
            return {"status": "failure", "reason": "unsupported_format"}

        # Extract text using python-docx
        from docx import Document

        doc = Document(full_path)

        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)

        # Simple table text extraction (tab-separated rows)
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                table_lines.append("\t".join(cells))
            if table_lines:
                parts.append("\n".join(table_lines))

        extracted = "\n\n".join(parts).strip()

        return {
            "status": "success",
            "result": extracted,
            "metadata": {
                "source_path": full_path,
                "format": "docx",
            },
        }

    except Exception:
        return {"status": "failure", "reason": "read_error"}
