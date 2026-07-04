"""Tests for read_docx tool.

SPRINT-11-SLICE-002 — Local document reader foundation.
"""

import os
import sys
import unittest

_PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
if _PROJECT_ROOT not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT)

from tools.read_docx import run


def _project_tmp(name: str) -> str:
    return os.path.join(_PROJECT_ROOT, "tmp", name)


class TestReadDocxTool(unittest.TestCase):

    def test_missing_file(self):
        """Reject missing file safely."""
        result = run("tmp/nonexistent_file.docx")
        self.assertEqual(result["status"], "failure")
        self.assertEqual(result["reason"], "file_not_found")

    def test_non_docx_extension(self):
        """Reject non-DOCX extension safely."""
        temp_path = _project_tmp("test_read_docx_wrong_ext.txt")
        os.makedirs(os.path.dirname(temp_path), exist_ok=True)
        with open(temp_path, "w", encoding="utf-8") as f:
            f.write("not a docx")
        try:
            result = run(temp_path)
            self.assertEqual(result["status"], "failure")
            self.assertEqual(result["reason"], "unsupported_format")
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def test_path_traversal_blocked(self):
        """Path outside base dir is blocked."""
        result = run("../../../etc/passwd.docx")
        self.assertEqual(result["status"], "failure")
        self.assertEqual(result["reason"], "path_safety_blocked")

    def test_extracts_paragraph_text(self):
        """Extract paragraph text from a real DOCX."""
        from docx import Document

        temp_path = _project_tmp("test_read_docx_paras.docx")
        os.makedirs(os.path.dirname(temp_path), exist_ok=True)
        try:
            doc = Document()
            doc.add_paragraph("First paragraph.")
            doc.add_paragraph("Second paragraph.")
            doc.save(temp_path)

            result = run(temp_path)
            self.assertEqual(result["status"], "success")
            self.assertIn("First paragraph.", result["result"])
            self.assertIn("Second paragraph.", result["result"])
            self.assertEqual(result["metadata"]["format"], "docx")
            self.assertIn("source_path", result["metadata"])
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def test_extracts_table_text(self):
        """Extract simple table text from a real DOCX."""
        from docx import Document

        temp_path = _project_tmp("test_read_docx_table.docx")
        os.makedirs(os.path.dirname(temp_path), exist_ok=True)
        try:
            doc = Document()
            doc.add_paragraph("Intro paragraph.")
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Header1"
            table.rows[0].cells[1].text = "Header2"
            table.rows[1].cells[0].text = "Cell1"
            table.rows[1].cells[1].text = "Cell2"
            doc.save(temp_path)

            result = run(temp_path)
            self.assertEqual(result["status"], "success")
            self.assertIn("Intro paragraph.", result["result"])
            self.assertIn("Header1", result["result"])
            self.assertIn("Cell2", result["result"])
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def test_invalid_docx_returns_read_error(self):
        """Invalid or unparseable DOCX returns read_error."""
        temp_path = _project_tmp("test_read_docx_invalid.docx")
        os.makedirs(os.path.dirname(temp_path), exist_ok=True)
        with open(temp_path, "wb") as f:
            f.write(b"not a docx")
        try:
            result = run(temp_path)
            self.assertEqual(result["status"], "failure")
            self.assertEqual(result["reason"], "read_error")
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)


if __name__ == "__main__":
    unittest.main()
